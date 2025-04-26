import streamlit as st
from fraud_keyword_detector import FraudKeywordDetector
import docx
import fitz  # PyMuPDF
import re
import pandas as pd
import io
import os
from collections import Counter

st.set_page_config(page_title="Fraud Risk Analyzer", layout="wide")
st.title("🔍 Fraud Risk Analyzer")

detector = FraudKeywordDetector()

def read_txt(file):
    return file.read().decode("utf-8")

def read_docx(file):
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def read_pdf(file):
    text = ""
    pdf = fitz.open(stream=file.read(), filetype="pdf")
    for page in pdf:
        text += page.get_text()
    return text

def extract_text(file, file_type):
    if file_type == "txt":
        return read_txt(file)
    elif file_type == "docx":
        return read_docx(file)
    elif file_type == "pdf":
        return read_pdf(file)
    else:
        return ""

def highlight_html(text, keywords_colors):
    tokens = re.findall(r'\w+|\W+', text)
    highlighted = ''
    for token in tokens:
        clean_token = re.sub(r'[^\w]', '', token).lower()
        color = keywords_colors.get(clean_token)
        if color:
            highlighted += f'<mark style="background-color: #{color}; font-weight:bold;">{token.strip()}</mark>' + token[len(token.strip()):]
        else:
            highlighted += token
    return highlighted

def highlight_docx(text, keywords_colors):
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    para = doc.add_paragraph()
    tokens = re.findall(r'\w+|\W+', text)
    for token in tokens:
        clean_token = re.sub(r'[^\w]', '', token).lower()
        run = para.add_run(token)
        if clean_token in keywords_colors:
            highlight = run._element.get_or_add_rPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), keywords_colors[clean_token])
            highlight.append(shd)
    return doc

def count_keywords(text, selected_keywords):
    tokens = re.findall(r'\w+', text.lower())
    counter = Counter(tokens)
    keyword_counts = {keyword: counter[keyword.lower()] for keyword in selected_keywords if counter[keyword.lower()] > 0}
    return keyword_counts

uploaded_file = st.file_uploader("Upload Document (.txt, .docx, .pdf)", type=["txt", "docx", "pdf"])

if uploaded_file:
    file_type = uploaded_file.name.split(".")[-1]
    file_base_name = os.path.splitext(uploaded_file.name)[0]
    raw_text = extract_text(uploaded_file, file_type)

    st.sidebar.header("🔖 Select Risk Categories to Highlight")
    selected_categories = st.sidebar.multiselect("Choose Categories:", list(detector.keyword_categories.keys()), default=list(detector.keyword_categories.keys())[:3])

    if selected_categories:
        keywords_colors = detector.get_keywords_and_colors(selected_categories)

        st.subheader("📄 Document Preview with Highlights")
        highlighted_text = highlight_html(raw_text, keywords_colors)
        st.markdown(highlighted_text, unsafe_allow_html=True)

        # Word Document download
        st.subheader("📥 Download Highlighted Word Document")
        docx_file = highlight_docx(raw_text, keywords_colors)
        output = io.BytesIO()
        docx_file.save(output)
        output.seek(0)
        download_filename_docx = f"highlighted_{file_base_name}.docx"

        st.download_button(
            label=f"Download {download_filename_docx}",
            data=output,
            file_name=download_filename_docx,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # CSV Keyword report download
        st.subheader("📈 Download Keyword Occurrence Report")
        keyword_counts = count_keywords(raw_text, keywords_colors.keys())
        if keyword_counts:
            report_df = pd.DataFrame(list(keyword_counts.items()), columns=["Keyword", "Occurrences"])
            csv_output = report_df.to_csv(index=False).encode('utf-8')
            download_filename_csv = f"keyword_report_{file_base_name}.csv"

            st.download_button(
                label=f"Download {download_filename_csv}",
                data=csv_output,
                file_name=download_filename_csv,
                mime="text/csv"
            )
        else:
            st.info("No selected keywords were found in the document.")

        # Legend
        st.subheader("🗂️ Legend for Highlights")
        legend_html = "<table style='width:100%; border-collapse: collapse;'>"
        legend_html += "<tr><th style='border: 1px solid black; padding: 5px;'>Category</th><th style='border: 1px solid black; padding: 5px;'>Color</th><th style='border: 1px solid black; padding: 5px;'>Keywords</th></tr>"
        for cat, data in detector.keyword_categories.items():
            if cat in selected_categories:
                color_cell = f"<td style='border: 1px solid black; background-color: #{data['color']}; padding: 5px;'></td>"
                keywords_text = ", ".join(data["keywords"])
                legend_html += f"<tr><td style='border: 1px solid black; padding: 5px;'>{cat}</td>{color_cell}<td style='border: 1px solid black; padding: 5px;'>{keywords_text}</td></tr>"
        legend_html += "</table>"
        st.markdown(legend_html, unsafe_allow_html=True)

st.sidebar.markdown("---")
st.sidebar.info("Developed by Ajay Rallapalli | Spring 2025")
